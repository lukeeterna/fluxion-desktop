#!/usr/bin/env python3
"""
Genera i 4 template GDPR FLUXION (lead magnet landing).

Output in landing/assets/gdpr/:
  - informativa-privacy.docx       (Art. 13 GDPR — Informativa clienti, 1 pagina SHORT)
  - registro-trattamenti.xlsx      (Art. 30 GDPR — Registro titolare, 7 trattamenti pre-compilati)
  - consenso-art9-sanitario.pdf    (Art. 9 GDPR — Consenso esplicito dati salute)
  - guida-gdpr-pmi.html            (Checklist 15 step + riferimenti normativi)

Disclaimer legale standard incluso su tutti i template.
Riferimenti: GDPR 2016/679 + D.Lgs 196/2003 + provv. Garante 9091942 / FAQ 9047529.

Run:
  python3 scripts/generate_gdpr_templates.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY, TA_CENTER

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "landing" / "assets" / "gdpr"
OUT_DIR.mkdir(parents=True, exist_ok=True)

DISCLAIMER = (
    "Questo documento è un modello informativo fornito a titolo orientativo e non costituisce "
    "consulenza legale. Le PMI restano responsabili dei propri adempimenti GDPR. Per studi medici, "
    "dentali, fisioterapici si consiglia verifica con DPO o consulente legale specializzato. "
    "Riferimenti: Reg. UE 2016/679 (GDPR) | D.Lgs 196/2003 mod. D.Lgs 101/2018 | garanteprivacy.it"
)


# ─── 1. Informativa Privacy DOCX ────────────────────────────────────

def generate_informativa() -> Path:
    doc = Document()

    # Margini stretti (1 pagina fronte-retro)
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_heading("INFORMATIVA SUL TRATTAMENTO DEI DATI PERSONALI", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("(ai sensi dell'art. 13 del Reg. UE 2016/679 — GDPR)")
    sub_run.italic = True
    sub_run.font.size = Pt(9)

    doc.add_paragraph()

    doc.add_heading("1. Chi siamo (Titolare del trattamento)", level=1)
    doc.add_paragraph(
        "Denominazione: [NOME ATTIVITÀ]\n"
        "Indirizzo: [INDIRIZZO COMPLETO]\n"
        "Email: [EMAIL]    —    Telefono: [TELEFONO]\n"
        "Titolare: [NOME E COGNOME TITOLARE]"
    )

    doc.add_heading("2. Quali dati raccogliamo e perché", level=1)
    doc.add_paragraph(
        "Trattiamo i seguenti dati per le finalità indicate. Tutti i trattamenti hanno una base "
        "giuridica documentata e nessun dato viene ceduto a terzi per fini commerciali."
    )
    doc.add_paragraph(
        "• Dati anagrafici (nome, cognome, telefono, email): per gestire prenotazioni, "
        "comunicazioni di servizio, fatturazione. Base: art. 6(1)(b) GDPR — esecuzione del contratto."
    )
    doc.add_paragraph(
        "• Dati relativi alla salute (allergie, controindicazioni, anamnesi) [SE APPLICABILE]: "
        "per garantire la sicurezza dei trattamenti. Base: art. 9(2)(a) GDPR — consenso esplicito."
    )
    doc.add_paragraph(
        "• WhatsApp / SMS: per conferme appuntamento (base contrattuale) e, separatamente, "
        "per offerte e promozioni (base art. 6(1)(a) — consenso, revocabile in ogni momento)."
    )
    doc.add_paragraph(
        "• Eventuali chiamate gestite da assistente vocale automatico: la prenotazione è sempre "
        "verificata dall'operatore e l'audio non viene conservato dopo l'elaborazione."
    )

    doc.add_heading("3. Con chi condividiamo i tuoi dati", level=1)
    doc.add_paragraph(
        "I tuoi dati sono trattati esclusivamente da personale autorizzato del nostro studio/attività. "
        "Il gestionale FLUXION conserva i dati localmente sul nostro PC, senza cessione a terzi né cloud "
        "non autorizzato. Eventuali fornitori tecnici (es. provider WhatsApp) sono nominati responsabili "
        "del trattamento ai sensi dell'art. 28 GDPR."
    )

    doc.add_heading("4. Per quanto tempo conserviamo i dati", level=1)
    doc.add_paragraph(
        "• Dati anagrafici: 7 anni dalla cessazione del rapporto.\n"
        "• Documenti contabili e fiscali: 10 anni (obbligo D.P.R. 600/1973 e Codice Civile art. 2220).\n"
        "• Consenso a marketing: fino a revoca + ulteriori 12 mesi per finalità di prova."
    )

    doc.add_heading("5. I tuoi diritti", level=1)
    doc.add_paragraph(
        "In ogni momento puoi esercitare i diritti previsti dagli artt. 15-22 GDPR: accesso ai dati, "
        "rettifica, cancellazione, limitazione, portabilità, opposizione e revoca del consenso. "
        "Puoi anche proporre reclamo al Garante per la protezione dei dati personali "
        "(www.garanteprivacy.it). Per esercitare i tuoi diritti scrivi a [EMAIL]."
    )

    doc.add_heading("6. Aggiornamenti", level=1)
    doc.add_paragraph(
        "Questa informativa è aggiornata al [DATA]. Eventuali modifiche saranno comunicate "
        "tempestivamente attraverso i nostri canali."
    )

    doc.add_paragraph()
    p_disclaimer = doc.add_paragraph()
    p_disclaimer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p_disclaimer.add_run(DISCLAIMER)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    out = OUT_DIR / "informativa-privacy.docx"
    doc.save(out)
    return out


# ─── 2. Registro Trattamenti XLSX ───────────────────────────────────

def generate_registro() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Registro trattamenti"

    headers = [
        "Nome trattamento",
        "Finalità",
        "Base giuridica (GDPR art.)",
        "Categorie interessati",
        "Categorie dati",
        "Dati art. 9 (sì/no)",
        "Destinatari",
        "Trasferimenti extra-UE",
        "Periodo conservazione",
        "Misure sicurezza",
        "Data ultima verifica",
    ]
    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    border = Border(
        left=Side(style="thin", color="999999"),
        right=Side(style="thin", color="999999"),
        top=Side(style="thin", color="999999"),
        bottom=Side(style="thin", color="999999"),
    )

    for col_idx, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border

    rows = [
        [
            "Anagrafica clienti",
            "Gestione clienti, prenotazioni, comunicazioni di servizio",
            "Art. 6(1)(b) — esecuzione del contratto",
            "Clienti dell'attività",
            "Nome, cognome, telefono, email, codice fiscale",
            "No",
            "Personale interno autorizzato",
            "No",
            "7 anni dalla cessazione",
            "Backup criptato locale, accesso con credenziali",
            "[DATA]",
        ],
        [
            "Prenotazioni e agenda",
            "Programmazione appuntamenti, reminder",
            "Art. 6(1)(b) — esecuzione del contratto",
            "Clienti",
            "Data, ora, servizio richiesto",
            "No",
            "Personale interno",
            "No (gestionale on-premise)",
            "3 anni",
            "Cifratura SQLite a riposo (AES-256)",
            "[DATA]",
        ],
        [
            "WhatsApp/SMS — conferme appuntamento",
            "Notifiche di servizio: conferma e reminder appuntamento",
            "Art. 6(1)(b) — esecuzione contratto",
            "Clienti",
            "Numero telefono, contenuto messaggio",
            "No",
            "Meta Platforms Inc. (provider WhatsApp Business)",
            "Sì — USA con SCCs (Standard Contractual Clauses)",
            "12 mesi",
            "Crittografia end-to-end WhatsApp",
            "[DATA]",
        ],
        [
            "Marketing e promozioni (WhatsApp/email)",
            "Invio offerte, promozioni, programmi fedeltà",
            "Art. 6(1)(a) — consenso esplicito",
            "Clienti che hanno dato consenso",
            "Numero telefono, email, storico acquisti",
            "No",
            "Meta Platforms / Provider email",
            "Sì — USA con SCCs",
            "Fino a revoca + 12 mesi",
            "Consenso registrato e revocabile",
            "[DATA]",
        ],
        [
            "Fatturazione elettronica",
            "Adempimenti fiscali e contabili",
            "Art. 6(1)(c) — obbligo legale",
            "Clienti, fornitori",
            "Dati anagrafici, IVA, codice destinatario, importi",
            "No",
            "Agenzia Entrate (SDI), commercialista",
            "No",
            "10 anni (art. 2220 Codice Civile)",
            "Trasmissione SDI cifrata, conservazione provider certificato",
            "[DATA]",
        ],
        [
            "Assistente vocale automatico (booking)",
            "Risposta a chiamate per prenotazioni quando l'attività è chiusa o occupata",
            "Art. 6(1)(b) — esecuzione contratto",
            "Chiamanti",
            "Audio chiamata (non conservato), trascrizione testo, dati prenotazione",
            "No",
            "Provider STT/LLM internazionali certificati",
            "Sì — USA con SCCs (DPA con provider)",
            "Trascrizione: 30 giorni. Audio: NON conservato.",
            "Disclosure obbligatoria all'inizio chiamata. NON profilazione (art. 22 NA — operatore conferma sempre).",
            "[DATA]",
        ],
        [
            "Dati salute [SOLO verticali sanitari/estetici]",
            "Anamnesi, allergie, controindicazioni per sicurezza trattamento",
            "Art. 9(2)(h) [sanitario professionale] OPPURE Art. 9(2)(a) [estetico/sportivo: consenso esplicito]",
            "Clienti/pazienti",
            "Allergie, patologie, farmaci, gravidanza, controindicazioni cutanee",
            "Sì — categorie particolari art. 9",
            "Solo professionista e collaboratori autorizzati",
            "No",
            "10 anni (cartella clinica) o cessazione rapporto",
            "Cifratura locale, accesso ristretto con audit log",
            "[DATA]",
        ],
    ]

    for row_idx, row in enumerate(rows, start=2):
        for col_idx, val in enumerate(row, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border
            cell.font = Font(name="Calibri", size=10)

    widths = [22, 32, 28, 22, 30, 14, 28, 24, 22, 32, 14]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    for row_idx in range(2, 2 + len(rows)):
        ws.row_dimensions[row_idx].height = 90

    ws.row_dimensions[1].height = 50

    # Sheet 2: istruzioni
    ws2 = wb.create_sheet("Istruzioni")
    ws2["A1"] = "ISTRUZIONI USO REGISTRO TRATTAMENTI"
    ws2["A1"].font = Font(name="Calibri", size=14, bold=True, color="1F4E78")
    instructions = [
        "",
        "1. Sostituisci [DATA] e [NOME ATTIVITÀ] con i tuoi dati reali.",
        "2. Rimuovi le righe non applicabili (es. 'Dati salute' se non gestisci sanitario/estetico).",
        "3. Aggiungi righe per altri trattamenti specifici (es. videosorveglianza, dipendenti).",
        "4. Aggiorna la colonna 'Data ultima verifica' almeno una volta l'anno (art. 30(3) GDPR).",
        "5. Conserva il registro in formato cartaceo o digitale, accessibile in caso di ispezione.",
        "",
        "OBBLIGO REGISTRO (art. 30 GDPR + Garante FAQ doc. 9047529):",
        "Tutte le PMI con trattamenti non occasionali (clienti ricorrenti) sono obbligate al registro,",
        "anche sotto i 250 dipendenti. L'agenda clienti di un salone, palestra, studio = trattamento",
        "non occasionale → obbligo del registro.",
        "",
        "ART. 9 GDPR (categorie particolari):",
        "Se tratti dati salute (anche solo allergie del cliente) ti serve consenso esplicito separato",
        "(art. 9(2)(a)) — modulo consenso-art9-sanitario.pdf incluso nel pacchetto FLUXION.",
        "I professionisti sanitari (medici, dentisti, fisio) possono operare in art. 9(2)(h) senza",
        "consenso separato per le finalità di cura, ma serve consenso per finalità ulteriori.",
        "",
        DISCLAIMER,
    ]
    for i, line in enumerate(instructions, start=2):
        c = ws2.cell(row=i, column=1, value=line)
        c.alignment = Alignment(wrap_text=True, vertical="top")
        c.font = Font(name="Calibri", size=11)
    ws2.column_dimensions["A"].width = 110

    out = OUT_DIR / "registro-trattamenti.xlsx"
    wb.save(out)
    return out


# ─── 3. Consenso Art. 9 Sanitario PDF ───────────────────────────────

def generate_consenso() -> Path:
    out = OUT_DIR / "consenso-art9-sanitario.pdf"
    doc = SimpleDocTemplate(
        str(out),
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2.2 * cm,
        rightMargin=2.2 * cm,
        title="Consenso al trattamento dei dati di cui all'art. 9 GDPR",
        author="FLUXION",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        fontSize=14,
        leading=18,
        alignment=TA_CENTER,
        spaceAfter=6,
        textColor=HexColor("#1F4E78"),
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontSize=10,
        leading=12,
        alignment=TA_CENTER,
        textColor=HexColor("#666666"),
        fontName="Helvetica-Oblique",
    )
    h2_style = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontSize=11,
        leading=14,
        spaceBefore=10,
        spaceAfter=4,
        textColor=HexColor("#1F4E78"),
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        leading=13,
        alignment=TA_JUSTIFY,
        spaceAfter=4,
    )
    field_style = ParagraphStyle(
        "Field",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=2,
    )
    disclaimer_style = ParagraphStyle(
        "Disclaimer",
        parent=styles["Normal"],
        fontSize=8,
        leading=10,
        alignment=TA_JUSTIFY,
        textColor=HexColor("#888888"),
        fontName="Helvetica-Oblique",
        spaceBefore=12,
    )

    story = []

    story.append(Paragraph("CONSENSO AL TRATTAMENTO DEI DATI", title_style))
    story.append(Paragraph(
        "di cui all'art. 9 del Reg. UE 2016/679 (GDPR) — categorie particolari",
        subtitle_style,
    ))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Titolare del trattamento", h2_style))
    story.append(Paragraph("Denominazione: ____________________________________________", field_style))
    story.append(Paragraph("Indirizzo: __________________________________________________", field_style))
    story.append(Paragraph("Email: ________________________ &nbsp;&nbsp; Tel: ___________________", field_style))

    story.append(Paragraph("Dati dell'interessato", h2_style))
    story.append(Paragraph("Nome e cognome: __________________________________________", field_style))
    story.append(Paragraph("Data e luogo di nascita: ___________________________________", field_style))
    story.append(Paragraph("Codice fiscale: ___________________________________________", field_style))

    story.append(Paragraph("Finalità del trattamento", h2_style))
    story.append(Paragraph(
        "Il Titolare tratterà i dati relativi alla salute (allergie, patologie pregresse, "
        "controindicazioni, anamnesi, eventuali farmaci in uso, gravidanza in corso, dati biometrici "
        "se rilevanti) esclusivamente per le seguenti finalità:",
        body_style,
    ))
    story.append(Paragraph(
        "• garantire la sicurezza del trattamento estetico/sportivo/sanitario;<br/>"
        "• personalizzare il servizio in base alle condizioni dell'interessato;<br/>"
        "• prevenire reazioni avverse;<br/>"
        "• adempiere a eventuali obblighi normativi.",
        body_style,
    ))

    story.append(Paragraph("Categorie di dati raccolti", h2_style))
    checklist = [
        ("[ ]", "Allergie e intolleranze"),
        ("[ ]", "Patologie pregresse o in corso"),
        ("[ ]", "Farmaci attualmente in uso"),
        ("[ ]", "Stato di gravidanza"),
        ("[ ]", "Interventi chirurgici recenti"),
        ("[ ]", "Controindicazioni cutanee/dermatologiche"),
        ("[ ]", "Altre condizioni rilevanti: __________________________"),
    ]
    for box, label in checklist:
        story.append(Paragraph(f"{box} &nbsp; {label}", field_style))

    story.append(Paragraph("Conservazione, diritti, base giuridica", h2_style))
    story.append(Paragraph(
        "I dati saranno conservati per il tempo necessario alle finalità (di norma 10 anni per la "
        "cartella clinica, secondo la prassi del settore). L'interessato può esercitare i diritti di "
        "accesso, rettifica, cancellazione, limitazione, portabilità, opposizione (artt. 15-22 GDPR) "
        "scrivendo al Titolare. Il presente consenso costituisce base giuridica ai sensi dell'art. 9(2)(a) "
        "GDPR per il trattamento delle categorie particolari di dati. Il consenso è libero, specifico "
        "e revocabile in ogni momento senza pregiudizio del trattamento svolto in precedenza.",
        body_style,
    ))

    story.append(Spacer(1, 8))
    story.append(Paragraph("Dichiarazione di consenso", h2_style))
    story.append(Paragraph(
        "Letta l'informativa ricevuta in forma estesa, l'interessato:",
        body_style,
    ))
    story.append(Spacer(1, 4))

    consent_data = [
        ["[ ] PRESTO il consenso", "[ ] NON presto il consenso"],
    ]
    consent_table = Table(consent_data, colWidths=[8 * cm, 8 * cm])
    consent_table.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 11),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOX", (0, 0), (-1, -1), 0.5, HexColor("#999999")),
        ("INNERGRID", (0, 0), (-1, -1), 0.3, HexColor("#cccccc")),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))
    story.append(consent_table)

    story.append(Spacer(1, 14))
    story.append(Paragraph("Luogo e data: ________________________________________", field_style))
    story.append(Spacer(1, 8))
    story.append(Paragraph("Firma dell'interessato: _______________________________", field_style))

    story.append(PageBreak())
    story.append(Paragraph("REVOCA DEL CONSENSO", h2_style))
    story.append(Paragraph(
        "Il sottoscritto revoca il consenso prestato in data _______________ "
        "al trattamento dei dati di cui all'art. 9 GDPR. La revoca non pregiudica la liceità del "
        "trattamento svolto in precedenza.",
        body_style,
    ))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Luogo e data: ________________________________________", field_style))
    story.append(Spacer(1, 8))
    story.append(Paragraph("Firma: ______________________________________________", field_style))

    story.append(Paragraph(DISCLAIMER, disclaimer_style))

    doc.build(story)
    return out


# ─── 4. Guida GDPR PMI HTML ─────────────────────────────────────────

def generate_guida() -> Path:
    out = OUT_DIR / "guida-gdpr-pmi.html"

    steps = [
        ("Sei titolare del trattamento?",
         "Se decidi finalità e modalità del trattamento (es. gestisci un'agenda clienti), sei titolare. "
         "Anche se sei una ditta individuale.",
         "Art. 4(7) GDPR"),
        ("Devi tenere il registro dei trattamenti?",
         "Sì, quasi sempre. Anche sotto 250 dipendenti l'obbligo scatta se il trattamento è non "
         "occasionale (l'agenda clienti ricorrenti = non occasionale) oppure se gestisci dati salute. "
         "Tutti i settori FLUXION rientrano.",
         "Art. 30 GDPR + Garante FAQ 9047529"),
        ("Hai un'informativa pronta da consegnare ai clienti?",
         "Obbligatoria al momento della prima raccolta dati. Deve essere concisa, chiara, in italiano "
         "semplice. Modello incluso nel pacchetto FLUXION.",
         "Art. 13 GDPR + Garante 9091942"),
        ("Devi nominare un DPO (Data Protection Officer)?",
         "Per PMI 1-15 dipendenti con operatività locale di norma NO. Diventa obbligatorio se l'attività "
         "principale è monitoraggio sistematico su larga scala o trattamento su larga scala di dati "
         "art. 9 (raro per la singola PMI).",
         "Art. 37 GDPR"),
        ("Tratti dati relativi alla salute? Hai il consenso esplicito?",
         "Allergie del cliente, controindicazioni, anamnesi = dati salute. Servono consenso esplicito "
         "scritto (art. 9(2)(a)) per estetisti, palestre, nail artist. Per medici/dentisti/fisio in "
         "operazioni di cura vale art. 9(2)(h) senza consenso separato, ma serve per finalità ulteriori.",
         "Art. 9 GDPR"),
        ("Hai un sito web? È presente cookie banner conforme?",
         "Banner senza pulsanti pre-spuntati, accetta/rifiuta sullo stesso livello, possibilità di "
         "configurare i cookie tecnici/profilazione. Le linee guida Garante 2021 sono ancora valide.",
         "Linee guida cookie Garante 2021 + D.Lgs 196/2003 art. 122"),
        ("Mandi messaggi WhatsApp di marketing senza consenso?",
         "Il fatto che il cliente ti abbia dato il numero per prenotare NON autorizza l'invio di "
         "promozioni. Serve consenso specifico per il marketing, registrato e revocabile.",
         "Art. 6(1)(a) GDPR + Garante prov. 6/3/2024"),
        ("Usi un assistente vocale AI al telefono? Avvisi il chiamante?",
         "Sì, sempre. Disclosure obbligatoria all'inizio della chiamata: \"Le ricordo che la chiamata "
         "è gestita da un assistente automatico, le sue richieste saranno elaborate per finalità di "
         "prenotazione\". L'AI di FLUXION (Sara) non rientra in art. 22 perché la prenotazione è "
         "sempre confermata dall'operatore.",
         "Art. 13 + Art. 22 GDPR"),
        ("Sai cosa fare in caso di data breach (PC rubato, ransomware)?",
         "Hai 72 ore per notificare al Garante via portale online. Se il rischio per gli interessati "
         "è alto, devi anche avvisare i clienti coinvolti. Tieni traccia di tutti gli eventi di breach "
         "in un registro dedicato.",
         "Art. 33 + Art. 34 GDPR"),
        ("Hai backup quotidiano dei dati attivo e funzionante?",
         "Misura di sicurezza richiesta dall'art. 32 GDPR. Backup criptato, almeno giornaliero, "
         "su supporto separato. Verifica periodicamente che il restore funzioni.",
         "Art. 32 GDPR — misure adeguate"),
        ("Dipendenti formati e nominati come autorizzati al trattamento?",
         "Chi accede ai dati clienti deve essere autorizzato per iscritto e formato sulle regole. "
         "Una formazione l'anno è il minimo. Tieni registro delle nomine.",
         "Art. 29 + Art. 32(4) GDPR"),
        ("Hai una procedura per gestire le richieste degli interessati?",
         "Cliente chiede i suoi dati, vuole cancellarli, vuole portarli a un altro fornitore — devi "
         "rispondere entro 30 giorni (1 mese). Procedura scritta + canale dedicato (email a [EMAIL]).",
         "Artt. 15-22 GDPR"),
        ("Usi provider USA (Google, Meta, Stripe, cloud)? Hai verificato le SCCs?",
         "Trasferimento dati extra-UE leggi Standard Contractual Clauses pubblicate dalla Commissione. "
         "Verifica che il provider abbia DPA aggiornato post-Schrems II.",
         "Artt. 44-49 GDPR + sentenza Schrems II"),
        ("Cancelli o anonimizzi dati dopo il termine di conservazione?",
         "Conservare \"per sicurezza\" oltre il necessario è violazione del principio di minimizzazione. "
         "Per dati clienti non più attivi: 7 anni di norma; per fatturazione: 10 anni.",
         "Art. 5(1)(e) GDPR"),
        ("Hai aggiornato il registro dei trattamenti nell'ultimo anno?",
         "Il registro va aggiornato in caso di nuovi trattamenti (es. attivazione voice agent, nuovo "
         "canale di marketing) e comunque almeno annualmente.",
         "Art. 30(3) GDPR"),
    ]

    items_html = ""
    for i, (q, body, ref) in enumerate(steps, start=1):
        items_html += f'''
        <li class="step">
          <label>
            <input type="checkbox" id="s{i}">
            <span class="num">{i}</span>
            <span class="content">
              <span class="q">{q}</span>
              <span class="body">{body}</span>
              <span class="ref">{ref}</span>
            </span>
          </label>
        </li>'''

    html = f"""<!DOCTYPE html>
<html lang="it">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Guida GDPR per PMI italiane — 15 step pratici</title>
<style>
  *{{box-sizing:border-box}}
  body{{margin:0;padding:0;font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,Helvetica,Arial,sans-serif;background:#f7f8fa;color:#1a1a1a;line-height:1.55}}
  .wrap{{max-width:780px;margin:0 auto;padding:32px 20px 80px}}
  header{{background:linear-gradient(135deg,#1F4E78,#2563eb);color:#fff;padding:36px 28px;border-radius:14px;margin-bottom:28px;box-shadow:0 6px 20px rgba(37,99,235,0.18)}}
  header h1{{margin:0 0 8px;font-size:26px;letter-spacing:-0.3px}}
  header p{{margin:0;color:#dbeafe;font-size:15px}}
  header .meta{{margin-top:14px;font-size:13px;color:#bfdbfe}}
  .intro{{background:#fff;border:1px solid #e4e7eb;border-radius:12px;padding:22px 26px;margin-bottom:24px;font-size:14px;color:#444}}
  .intro strong{{color:#1F4E78}}
  ol.checklist{{list-style:none;padding:0;margin:0;counter-reset:none}}
  li.step{{background:#fff;border:1px solid #e4e7eb;border-radius:10px;margin-bottom:12px;transition:border-color .15s}}
  li.step:hover{{border-color:#2563eb}}
  li.step label{{display:flex;align-items:flex-start;padding:18px 22px;cursor:pointer;gap:14px}}
  li.step input[type="checkbox"]{{margin-top:4px;width:20px;height:20px;cursor:pointer;accent-color:#2563eb;flex-shrink:0}}
  li.step .num{{font-weight:700;color:#2563eb;min-width:28px;font-size:15px;flex-shrink:0}}
  li.step .content{{display:flex;flex-direction:column;gap:6px;flex:1}}
  li.step .q{{font-weight:600;font-size:15.5px;color:#1a1a1a}}
  li.step .body{{font-size:14px;color:#555;line-height:1.55}}
  li.step .ref{{font-size:12px;color:#2563eb;font-weight:500;margin-top:2px}}
  li.step input:checked ~ .content .q{{text-decoration:line-through;color:#999}}
  .disclaimer{{margin-top:36px;padding:18px 22px;background:#fff8e1;border:1px solid #f5d486;border-radius:10px;font-size:12px;color:#7a5d00;line-height:1.6}}
  .footer{{margin-top:30px;text-align:center;font-size:12px;color:#888}}
  .footer a{{color:#2563eb;text-decoration:none}}
  @media print{{
    body{{background:#fff}}
    .wrap{{padding:8px}}
    header{{box-shadow:none;background:#1F4E78}}
    li.step{{break-inside:avoid;page-break-inside:avoid}}
  }}
</style>
</head>
<body>
<div class="wrap">
  <header>
    <h1>Guida GDPR per PMI italiane</h1>
    <p>15 step pratici per essere a posto con il regolamento — checklist interattiva.</p>
    <div class="meta">FLUXION — Risorse GDPR gratuite</div>
  </header>

  <div class="intro">
    Questa checklist copre i punti che il <strong>Garante per la protezione dei dati personali</strong>
    verifica con maggior frequenza nelle ispezioni a saloni, palestre, studi medici, centri estetici,
    officine, attività con clienti ricorrenti.
    <br><br>
    Spunta ogni step quando hai chiuso il punto. Stampa il documento (Cmd/Ctrl + P) per la riunione
    interna. I riferimenti normativi sono linkabili al testo ufficiale GDPR su EUR-Lex.
  </div>

  <ol class="checklist">{items_html}
  </ol>

  <div class="disclaimer">
    <strong>Avvertenza</strong> — {DISCLAIMER}
  </div>

  <div class="footer">
    Versione 1.0 — Riferimenti ufficiali: <a href="https://www.garanteprivacy.it/" target="_blank" rel="noopener">garanteprivacy.it</a> ·
    <a href="https://eur-lex.europa.eu/legal-content/IT/TXT/HTML/?uri=CELEX:32016R0679" target="_blank" rel="noopener">GDPR su EUR-Lex</a>
  </div>
</div>
</body>
</html>
"""
    out.write_text(html, encoding="utf-8")
    return out


# ─── Main ───────────────────────────────────────────────────────────

def main() -> None:
    print("Generating GDPR templates → landing/assets/gdpr/")
    paths = [
        generate_informativa(),
        generate_registro(),
        generate_consenso(),
        generate_guida(),
    ]
    for p in paths:
        size_kb = p.stat().st_size / 1024
        print(f"  OK  {p.name}  ({size_kb:.1f} KB)")
    print(f"Total: {len(paths)} files in {OUT_DIR}")


if __name__ == "__main__":
    main()
